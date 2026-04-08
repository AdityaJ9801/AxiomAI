#!/usr/bin/env python3
"""
Report Editor MCP Tool
Handles report generation and editing operations
"""

import os
from datetime import datetime
from typing import Dict, Any, List, Optional
from .base_mcp_tool import BaseMCPTool

class ReportEditor(BaseMCPTool):
    """MCP tool for report editing operations"""
    
    def __init__(self):
        super().__init__("report_editor")
        self.report_data = {
            "title": "",
            "author": "",
            "created_date": "",
            "sections": []
        }
        
    def get_capabilities(self) -> Dict[str, Any]:
        """Return tool capabilities"""
        return {
            "name": "report_editor",
            "description": "Report generation and editing operations",
            "actions": [
                "create_new_report",
                "add_section",
                "update_section",
                "delete_section",
                "add_analysis_results",
                "add_visualization",
                "export_to_word",
                "export_to_pdf",
                "get_report_data"
            ]
        }
    
    def execute(self, action: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """Execute report editing action"""
        try:
            if action == "create_new_report":
                return self.create_new_report(
                    parameters.get("title", "Untitled Report"),
                    parameters.get("author", "Unknown Author")
                )
            elif action == "add_section":
                return self.add_section(
                    parameters.get("heading", "Untitled Section"),
                    parameters.get("content", ""),
                    parameters.get("section_type", "text")
                )
            elif action == "update_section":
                return self.update_section(
                    parameters.get("section_id"),
                    parameters.get("heading"),
                    parameters.get("content")
                )
            elif action == "delete_section":
                return self.delete_section(parameters.get("section_id"))
            elif action == "add_analysis_results":
                return self.add_analysis_results(
                    parameters.get("analysis_type"),
                    parameters.get("results")
                )
            elif action == "add_visualization":
                return self.add_visualization(
                    parameters.get("chart_type"),
                    parameters.get("data"),
                    parameters.get("title", "Chart")
                )
            elif action == "export_to_word":
                return self.export_to_word(parameters.get("filename", "report.docx"))
            elif action == "export_to_pdf":
                return self.export_to_pdf(parameters.get("filename", "report.pdf"))
            elif action == "get_report_data":
                return self.get_report_data()
            else:
                return {"success": False, "error": f"Unknown action: {action}"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def create_new_report(self, title: str, author: str) -> str:
        """Create a new report"""
        self.report_data = {
            "title": title,
            "author": author,
            "created_date": datetime.now().isoformat(),
            "sections": []
        }
        
        return f"New report '{title}' created successfully"
    
    def add_section(self, heading: str, content: str, section_type: str = "text") -> Dict[str, Any]:
        """Add a new section to the report"""
        section_id = len(self.report_data["sections"])
        
        section = {
            "id": section_id,
            "heading": heading,
            "content": content,
            "type": section_type,
            "created_date": datetime.now().isoformat()
        }
        
        self.report_data["sections"].append(section)
        
        return {
            "success": True,
            "section_id": section_id,
            "message": f"Section '{heading}' added successfully"
        }
    
    def update_section(self, section_id: int, heading: str = None, content: str = None) -> Dict[str, Any]:
        """Update an existing section"""
        if section_id >= len(self.report_data["sections"]):
            return {"success": False, "error": "Section not found"}
        
        section = self.report_data["sections"][section_id]
        
        if heading is not None:
            section["heading"] = heading
        if content is not None:
            section["content"] = content
        
        section["modified_date"] = datetime.now().isoformat()
        
        return {
            "success": True,
            "message": f"Section {section_id} updated successfully"
        }
    
    def delete_section(self, section_id: int) -> Dict[str, Any]:
        """Delete a section"""
        if section_id >= len(self.report_data["sections"]):
            return {"success": False, "error": "Section not found"}
        
        deleted_section = self.report_data["sections"].pop(section_id)
        
        # Update IDs of remaining sections
        for i, section in enumerate(self.report_data["sections"]):
            section["id"] = i
        
        return {
            "success": True,
            "message": f"Section '{deleted_section['heading']}' deleted successfully"
        }
    
    def add_analysis_results(self, analysis_type: str, results: Dict[str, Any]) -> Dict[str, Any]:
        """Add analysis results as a section"""
        content = self._format_analysis_results(analysis_type, results)
        
        return self.add_section(
            heading=f"{analysis_type.title()} Analysis Results",
            content=content,
            section_type="analysis"
        )
    
    def _format_analysis_results(self, analysis_type: str, results: Dict[str, Any]) -> str:
        """Format analysis results for report"""
        if analysis_type == "descriptive":
            content = "## Descriptive Statistics\n\n"
            
            if "statistics" in results:
                content += "### Summary Statistics\n"
                for column, stats in results["statistics"].items():
                    content += f"\n**{column}:**\n"
                    for stat, value in stats.items():
                        content += f"- {stat}: {value:.4f}\n"
            
            if "correlations" in results and results["correlations"]:
                content += "\n### Correlation Analysis\n"
                content += "Strong correlations found between variables.\n"
        
        elif analysis_type == "regression":
            content = "## Regression Analysis Results\n\n"
            
            if "r_squared" in results:
                content += f"**R-squared:** {results['r_squared']:.4f}\n\n"
            
            if "coefficients" in results:
                content += "### Coefficients\n"
                for var, coef in results["coefficients"].items():
                    content += f"- {var}: {coef:.4f}\n"
        
        elif analysis_type == "timeseries":
            content = "## Time Series Analysis Results\n\n"
            
            if "trend" in results:
                content += f"**Trend:** {results['trend']}\n\n"
            
            if "forecast" in results:
                content += "### Forecast\n"
                content += f"Generated forecast for {results['forecast'].get('periods', 0)} periods.\n"
        
        else:
            content = f"## {analysis_type.title()} Analysis\n\nResults: {str(results)}"
        
        return content
    
    def add_visualization(self, chart_type: str, data: Dict[str, Any], title: str = "Chart") -> Dict[str, Any]:
        """Add a visualization section"""
        content = f"## {title}\n\n"
        content += f"Chart Type: {chart_type}\n\n"
        content += "![Chart](chart_placeholder.png)\n\n"
        content += f"Data points: {len(data.get('data', []))}\n"
        
        return self.add_section(
            heading=title,
            content=content,
            section_type="visualization"
        )
    
    def export_to_word(self, filename: str) -> Dict[str, Any]:
        """Export report to Word document"""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Add title
            title = doc.add_heading(self.report_data["title"], 0)
            
            # Add metadata
            doc.add_paragraph(f"Author: {self.report_data['author']}")
            doc.add_paragraph(f"Created: {self.report_data['created_date']}")
            doc.add_paragraph("")
            
            # Add sections
            for section in self.report_data["sections"]:
                doc.add_heading(section["heading"], level=1)
                doc.add_paragraph(section["content"])
                doc.add_paragraph("")
            
            # Save document
            doc.save(filename)
            
            return {
                "success": True,
                "filename": filename,
                "message": f"Report exported to {filename}"
            }
        
        except ImportError:
            return {
                "success": False,
                "error": "python-docx library not installed. Install with: pip install python-docx"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def export_to_pdf(self, filename: str) -> Dict[str, Any]:
        """Export report to PDF"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(filename, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
            )
            story.append(Paragraph(self.report_data["title"], title_style))
            story.append(Spacer(1, 12))
            
            # Metadata
            story.append(Paragraph(f"Author: {self.report_data['author']}", styles['Normal']))
            story.append(Paragraph(f"Created: {self.report_data['created_date']}", styles['Normal']))
            story.append(Spacer(1, 24))
            
            # Sections
            for section in self.report_data["sections"]:
                story.append(Paragraph(section["heading"], styles['Heading2']))
                story.append(Spacer(1, 12))
                
                # Split content by lines and add as paragraphs
                for line in section["content"].split('\n'):
                    if line.strip():
                        story.append(Paragraph(line, styles['Normal']))
                
                story.append(Spacer(1, 18))
            
            doc.build(story)
            
            return {
                "success": True,
                "filename": filename,
                "message": f"Report exported to {filename}"
            }
        
        except ImportError:
            return {
                "success": False,
                "error": "reportlab library not installed. Install with: pip install reportlab"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def get_report_data(self) -> Dict[str, Any]:
        """Get current report data"""
        return {
            "success": True,
            "report": self.report_data,
            "section_count": len(self.report_data["sections"]),
            "word_count": sum(len(section["content"].split()) for section in self.report_data["sections"])
        }